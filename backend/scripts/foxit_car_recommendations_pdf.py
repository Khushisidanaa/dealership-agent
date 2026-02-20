#!/usr/bin/env python3
"""
Generate a PDF of car recommendations using Foxit Document Generation API.

Converts a JSON list of car recommendations into a branded PDF report.

Setup:
  pip install python-docx requests

  Set env vars (or use backend/.env):
    FOXIT_CLIENT_ID=your-client-id
    FOXIT_CLIENT_SECRET=your-client-secret
    FOXIT_API_HOST=https://na1.fusion.foxit.com   # or your Foxit doc-gen host

Run:
  python backend/scripts/foxit_car_recommendations_pdf.py

Output:
  car_recommendations.pdf in the current directory.
"""

import base64
import io
import os
import sys
from pathlib import Path

# Add backend to path for dotenv
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parents[1] / ".env")
except ImportError:
    pass

try:
    import requests
except ImportError:
    print("Run: pip install requests")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Run: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Sample car recommendations (matches VehicleListingResult shape)
# ---------------------------------------------------------------------------

SAMPLE_CARS = [
    {
        "heading": "2023 Toyota Camry SE FWD",
        "price": 22995,
        "miles": 22473,
        "dealer_name": "Autosavvy Draper",
        "dealer_phone": "833-605-2526",
        "vin": "4T1G11AK3PU787980",
        "colors": {"exterior": "Gray", "interior": ""},
    },
    {
        "heading": "2023 Toyota Camry XSE",
        "price": 23995,
        "miles": 36200,
        "dealer_name": "Velocity Cars",
        "dealer_phone": "801-555-1234",
        "vin": "4T1K61AK0PU180818",
        "colors": {"exterior": "Silver", "interior": "Black"},
    },
    {
        "heading": "2023 Toyota Camry LE",
        "price": 20995,
        "miles": 45100,
        "dealer_name": "Ken Garff Honda",
        "dealer_phone": "801-555-5678",
        "vin": "4T1BF1FK7NU123456",
        "colors": {"exterior": "White", "interior": "Gray"},
    },
]


def create_template() -> bytes:
    """
    Create a Word template with Foxit tokens for car recommendations.
    Tokens: {{today}}, {{TableStart:cars}}, {{TableEnd:cars}},
    per-row: {{heading}}, {{price}}, {{miles}}, {{dealer_name}}
    """
    doc2 = Document()
    doc2.add_paragraph()
    p2 = doc2.add_paragraph()
    p2.add_run("Car Recommendations Report\n").bold = True
    p2.add_run("Generated on {{today}}\n\n")
    doc2.add_paragraph("Here are your personalized car recommendations:")
    doc2.add_paragraph()

    # Add TableStart token
    doc2.add_paragraph("{{TableStart:cars}}")

    table = doc2.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = "Vehicle", "Price", "Mileage", "Dealer"
    r = table.rows[1].cells
    r[0].text, r[1].text, r[2].text, r[3].text = (
        "{{heading}}",
        "{{price}}",
        "{{miles}}",
        "{{dealer_name}}",
    )

    doc2.add_paragraph("{{TableEnd:cars}}")
    doc2.add_paragraph()

    buf = io.BytesIO()
    doc2.save(buf)
    buf.seek(0)
    return buf.read()


def prepare_car_data(cars: list[dict]) -> list[dict]:
    """Flatten car objects for the template (price/miles as strings)."""
    result = []
    for c in cars:
        dealer = c.get("dealer") or {}
        dealer_name = c.get("dealer_name") or dealer.get("name", "")
        price_val = c.get("price")
        miles_val = c.get("miles")
        result.append({
            "heading": c.get("heading", c.get("title", "Unknown")),
            "price": f"${price_val:,}" if price_val is not None else "N/A",
            "miles": f"{miles_val:,} mi" if miles_val is not None else "N/A",
            "dealer_name": dealer_name,
        })
    return result


def generate_pdf(
    template_bytes: bytes,
    data: dict,
    client_id: str,
    client_secret: str,
    host: str,
) -> bytes:
    """Call Foxit Document Generation API; return PDF bytes."""
    url = f"{host.rstrip('/')}/document-generation/api/GenerateDocumentBase64"
    headers = {"client_id": client_id, "client_secret": client_secret}
    body = {
        "outputFormat": "pdf",
        "documentValues": data,
        "base64FileString": base64.b64encode(template_bytes).decode("utf-8"),
    }
    resp = requests.post(url, json=body, headers=headers, timeout=60)
    if not resp.ok:
        try:
            err_body = resp.json()
        except Exception:
            err_body = resp.text
        raise RuntimeError(
            f"Foxit API error {resp.status_code}: {resp.reason}\n{err_body}"
        )
    result = resp.json()

    if result.get("base64FileString") is None:
        raise RuntimeError(f"Foxit API error: {result}")

    return base64.b64decode(result["base64FileString"].encode("ascii"))


def main():
    client_id = os.environ.get("FOXIT_CLIENT_ID") or os.environ.get("CLIENT_ID")
    client_secret = (
        os.environ.get("FOXIT_CLIENT_SECRET") or os.environ.get("CLIENT_SECRET")
    )
    host = (
        os.environ.get("FOXIT_API_HOST")
        or os.environ.get("HOST")
        or "https://na1.fusion.foxit.com"
    )

    if not client_id or not client_secret:
        print(
            "Set FOXIT_CLIENT_ID and FOXIT_CLIENT_SECRET "
            "(or CLIENT_ID, CLIENT_SECRET)"
        )
        print("You can copy from backend/.env")
        sys.exit(1)

    cars = prepare_car_data(SAMPLE_CARS)
    data = {
        "title": "Your Car Recommendations",
        # Do not send "today" - Foxit fills {{today}} automatically; sending it causes duplicate-property error
        "cars": cars,
    }

    print("Creating Word template...")
    template = create_template()

    print("Calling Foxit Document Generation API...")
    pdf_bytes = generate_pdf(template, data, client_id, client_secret, host)

    output_path = Path.cwd() / "car_recommendations.pdf"
    output_path.write_bytes(pdf_bytes)
    print(f"Done! PDF saved to {output_path}")
    print(f"  ({len(pdf_bytes):,} bytes, {len(cars)} cars)")


if __name__ == "__main__":
    main()
