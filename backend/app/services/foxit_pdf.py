"""
Foxit Document Generation service for dashboard PDF export.

Generates a PDF report from dashboard data using the Foxit Document Generation API.
"""

import base64
import io
import logging
from typing import Any

import requests
from docx import Document

from app.config import get_settings

log = logging.getLogger(__name__)


def _create_dashboard_template() -> bytes:
    """
    Create a Word template with Foxit tokens for the dashboard report.
    Sections: Summary, Dealer Call Results, All Vehicles
    """
    doc = Document()

    # Title and metadata
    title = doc.add_paragraph()
    title_run = title.add_run("Vehicle Research Report")
    title_run.bold = True
    title_run.font.size = None  # use default
    title.alignment = 0  # LEFT

    sub = doc.add_paragraph()
    sub.add_run("Generated on {{today}}\n")
    sub.add_run("{{report_summary}}")

    doc.add_paragraph()

    # Dealer Call Results section
    sect = doc.add_paragraph()
    sect_run = sect.add_run("Dealer Call Results")
    sect_run.bold = True
    doc.add_paragraph("Vehicles we contacted — availability, verdict, and notes from dealer calls.")
    doc.add_paragraph()
    doc.add_paragraph("{{TableStart:call_results}}")
    table1 = doc.add_table(rows=2, cols=6)
    table1.style = "Table Grid"
    h = table1.rows[0].cells
    h[0].text, h[1].text, h[2].text = "Vehicle", "Price", "Dealer"
    h[3].text, h[4].text, h[5].text = "Availability", "Verdict", "Notes"
    r = table1.rows[1].cells
    r[0].text = "{{heading}}"
    r[1].text = "{{price}}"
    r[2].text = "{{dealer_name}}"
    r[3].text = "{{availability}}"
    r[4].text = "{{verdict}}"
    r[5].text = "{{notes}}"
    doc.add_paragraph("{{TableEnd:call_results}}")
    doc.add_paragraph()

    # All Vehicles section
    sect2 = doc.add_paragraph()
    sect2_run = sect2.add_run("All Vehicles")
    sect2_run.bold = True
    doc.add_paragraph("Your shortlisted vehicles.")
    doc.add_paragraph()
    doc.add_paragraph("{{TableStart:all_vehicles}}")
    table2 = doc.add_table(rows=2, cols=4)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text, h2[1].text, h2[2].text, h2[3].text = "Vehicle", "Price", "Mileage", "Dealer"
    r2 = table2.rows[1].cells
    r2[0].text = "{{heading}}"
    r2[1].text = "{{price}}"
    r2[2].text = "{{miles}}"
    r2[3].text = "{{dealer_name}}"
    doc.add_paragraph("{{TableEnd:all_vehicles}}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _format_availability(details: dict | None) -> str:
    if not details:
        return "Unknown"
    av = details.get("is_available")
    if av is True:
        return "Available"
    if av is False:
        return "Sold / Unavailable"
    return "Unknown"


def _format_verdict(details: dict | None) -> str:
    if not details:
        return ""
    return details.get("recommendation") or ""


def _format_notes(
    details: dict | None,
    summary: str | None,
    test_drive: str | None,
) -> str:
    parts = []
    if details and details.get("key_takeaways"):
        parts.append(str(details["key_takeaways"]))
    if summary and summary not in parts:
        parts.append(summary)
    if test_drive:
        parts.append(f"Test drive: {test_drive}")
    return " | ".join(parts) if parts else "-"


def prepare_dashboard_data(
    vehicles: list[dict],
    communication_status: list[dict],
    bookings_by_vehicle: dict[str, dict],
) -> dict[str, Any]:
    """Transform dashboard data for Foxit template."""
    comm_by_id = {c["vehicle_id"]: c for c in communication_status}
    contacted_ids = {
        vid
        for vid, c in comm_by_id.items()
        if c.get("call_made") or c.get("text_sent")
    }

    call_results = []
    all_vehicles = []

    for v in vehicles:
        vid = v.get("vehicle_id", "")
        comm = comm_by_id.get(vid) or {}
        details = comm.get("call_details") or {}
        summary = comm.get("response") or ""
        booking = bookings_by_vehicle.get(vid)
        test_drive = ""
        if booking:
            test_drive = f"{booking.get('scheduled_date', '')} at {booking.get('scheduled_time', '')}".strip()

        title = v.get("title") or v.get("heading") or "Unknown"
        price_val = v.get("price")
        price_str = f"${price_val:,.0f}" if price_val is not None else "N/A"
        miles_val = v.get("mileage") or v.get("miles")
        miles_str = f"{miles_val:,} mi" if miles_val is not None else "N/A"
        dealer = v.get("dealer_name") or ""

        row = {
            "heading": title,
            "price": price_str,
            "miles": miles_str,
            "dealer_name": dealer,
        }

        if vid in contacted_ids:
            call_results.append({
                **row,
                "availability": _format_availability(details),
                "verdict": _format_verdict(details),
                "notes": _format_notes(details, summary, test_drive or None),
            })
        else:
            all_vehicles.append(row)

    # Summary for report header
    total = len(call_results) + len(all_vehicles)
    if total == 0:
        report_summary = "No vehicles in this report."
    elif len(call_results) == 0:
        report_summary = f"{len(all_vehicles)} vehicle{'s' if len(all_vehicles) != 1 else ''} in your shortlist."
    elif len(all_vehicles) == 0:
        report_summary = f"{len(call_results)} vehicle{'s' if len(call_results) != 1 else ''} contacted. See call details below."
    else:
        report_summary = f"{total} vehicles total — {len(call_results)} contacted with call results, {len(all_vehicles)} in your shortlist."

    # When no call results, add one placeholder row so the table has content
    if not call_results and all_vehicles:
        call_results = [{
            "heading": "(No dealer calls yet)",
            "price": "-",
            "dealer_name": "-",
            "availability": "-",
            "verdict": "-",
            "notes": "Run the analyze flow to call dealers and populate this section.",
        }]
    elif not call_results:
        call_results = [{
            "heading": "(No data)",
            "price": "-",
            "dealer_name": "-",
            "availability": "-",
            "verdict": "-",
            "notes": "Add vehicles and run analyze to see call results.",
        }]

    return {
        "report_summary": report_summary,
        "call_results": call_results,
        "all_vehicles": all_vehicles,
    }


def generate_dashboard_pdf(data: dict[str, Any]) -> bytes:
    """Call Foxit Document Generation API; return PDF bytes."""
    settings = get_settings()
    client_id = settings.foxit_client_id
    client_secret = settings.foxit_client_secret
    host = settings.foxit_api_host.rstrip("/")

    if not client_id or not client_secret:
        raise ValueError(
            "Foxit credentials not configured. Set FOXIT_CLIENT_ID and FOXIT_CLIENT_SECRET in .env"
        )

    template_bytes = _create_dashboard_template()
    url = f"{host}/document-generation/api/GenerateDocumentBase64"
    headers = {"client_id": client_id, "client_secret": client_secret}
    body = {
        "outputFormat": "pdf",
        "documentValues": data,
        "base64FileString": base64.b64encode(template_bytes).decode("utf-8"),
    }
    resp = requests.post(url, json=body, headers=headers, timeout=60)
    if not resp.ok:
        try:
            err_body = resp.json()
        except Exception:
            err_body = resp.text
        raise RuntimeError(f"Foxit API error {resp.status_code}: {err_body}")

    result = resp.json()
    if result.get("base64FileString") is None:
        raise RuntimeError(f"Foxit API error: no PDF in response")

    return base64.b64decode(result["base64FileString"].encode("ascii"))
